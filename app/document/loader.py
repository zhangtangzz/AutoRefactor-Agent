import os
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document


class DocumentLoader:
    """统一的文档加载器，支持 PDF / Word / Excel / TXT"""

    @staticmethod
    def load_pdf(file_path: str) -> List[Document]:
        docs = []
        # 优先使用 pdfplumber（对中文支持更好）
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        docs.append(Document(
                            page_content=text.strip(),
                            metadata={"source": file_path, "page": i + 1, "type": "pdf"},
                        ))
        except Exception:
            pass

        if not docs:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        docs.append(Document(
                            page_content=text.strip(),
                            metadata={"source": file_path, "page": i + 1, "type": "pdf"},
                        ))
            except Exception as e:
                raise RuntimeError(f"无法解析PDF文件 {file_path}: {e}")

        return docs

    @staticmethod
    def load_docx(file_path: str) -> List[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        full_text = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )
        if not full_text:
            raise RuntimeError(f"Word文档无有效文本: {file_path}")
        return [Document(
            page_content=full_text,
            metadata={"source": file_path, "type": "docx"},
        )]

    @staticmethod
    def load_excel(file_path: str) -> List[Document]:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        docs = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                row_str = " | ".join(
                    str(cell) for cell in row if cell is not None
                )
                if row_str.strip():
                    rows_text.append(row_str)
            if rows_text:
                docs.append(Document(
                    page_content=f"工作表 [{sheet_name}]:\n" + "\n".join(rows_text),
                    metadata={"source": file_path, "sheet": sheet_name, "type": "excel"},
                ))
        wb.close()
        return docs

    @staticmethod
    def load_txt(file_path: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        if not text.strip():
            raise RuntimeError(f"文本文件无内容: {file_path}")
        return [Document(
            page_content=text,
            metadata={"source": file_path, "type": "txt"},
        )]

    @classmethod
    def load(cls, file_path: str) -> List[Document]:
        ext = Path(file_path).suffix.lower()
        loaders = {
            ".pdf": cls.load_pdf,
            ".docx": cls.load_docx,
            ".doc": cls.load_docx,
            ".xlsx": cls.load_excel,
            ".xls": cls.load_excel,
            ".txt": cls.load_txt,
            ".md": cls.load_txt,
            ".csv": cls.load_txt,
        }
        loader = loaders.get(ext)
        if loader is None:
            raise ValueError(f"不支持的文件格式: {ext}")
        return loader(file_path)

    @classmethod
    def load_directory(cls, directory: str) -> List[Document]:
        all_docs = []
        supported = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".md", ".csv"}
        for root, _, files in os.walk(directory):
            for fname in files:
                ext = Path(fname).suffix.lower()
                if ext in supported:
                    try:
                        docs = cls.load(os.path.join(root, fname))
                        all_docs.extend(docs)
                    except Exception as e:
                        print(f"[警告] 跳过 {fname}: {e}")
        return all_docs
